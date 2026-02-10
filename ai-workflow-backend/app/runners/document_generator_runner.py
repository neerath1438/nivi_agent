"""Document Generator Node Runner - Multi-format export."""
import os
from datetime import datetime
from typing import Dict, Any
from app.runners.base_runner import BaseRunner

# Import libraries
try:
    from fpdf import FPDF
except ImportError:
    FPDF = None

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None


class DocumentGeneratorRunner(BaseRunner):
    """Runner for generating documents in multiple formats."""
    
    def run(self, node_data: Dict[str, Any], state: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate document from content in various formats.
        
        Supported formats: PDF, DOCX, TXT, MD
        """
        # Get content from previous node
        content = None
        if "summary" in state and state["summary"]:
            content = str(state["summary"])
        elif "output" in state and state["output"]:
            content = str(state["output"])
        else:
            content = node_data.get("content", "")
        
        if not content or not content.strip():
            return {
                "error": "No content provided for document generation",
                "output": "⚠️ Please provide content to export"
            }
        
        # Get configuration
        format_type = node_data.get("format", "pdf").lower()
        title = node_data.get("title", "AI Generated Document")
        template = node_data.get("template", "simple")  # simple/professional/executive
        
        # Get metadata from state if available
        metadata = state.get("metadata", {})
        
        print(f"📄 Generating {format_type.upper()} document: {title}")
        
        # Route to appropriate generator
        if format_type == "pdf":
            return self._generate_pdf(content, title, template, metadata)
        elif format_type == "docx":
            return self._generate_docx(content, title, template, metadata)
        elif format_type == "txt":
            return self._generate_txt(content, title, metadata)
        elif format_type in ["md", "markdown"]:
            return self._generate_markdown(content, title, metadata)
        else:
            return {
                "error": f"Unsupported format: {format_type}",
                "output": f"❌ Format '{format_type}' not supported. Use: pdf, docx, txt, or md"
            }
    
    def _generate_pdf(
        self,
        content: str,
        title: str,
        template: str,
        metadata: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Generate PDF document."""
        if not FPDF:
            return {"error": "FPDF library not installed", "output": "❌ PDF Error: Library missing"}
        
        try:
            pdf = FPDF()
            pdf.add_page()
            
            # Professional template styling
            if template == "professional":
                # Header with border
                pdf.set_fill_color(66, 126, 234)
                pdf.rect(0, 0, 210, 30, 'F')
                pdf.set_text_color(255, 255, 255)
                pdf.set_font("Arial", "B", 20)
                pdf.cell(0, 30, title, ln=True, align="C")
                pdf.set_text_color(0, 0, 0)
                pdf.ln(5)
            elif template == "executive":
                # Executive header
                pdf.set_font("Arial", "B", 18)
                pdf.cell(0, 10, title, ln=True, align="L")
                pdf.set_draw_color(66, 126, 234)
                pdf.line(10, 20, 200, 20)
                pdf.ln(5)
            else:
                # Simple template
                pdf.set_font("Arial", "B", 16)
                pdf.cell(0, 10, title, ln=True, align="C")
            
            pdf.ln(5)
            
            # Metadata section
            # Use a Unicode-capable font if possible
            # Standard PDF fonts (Arial, Helvetica, etc.) don't support Tamil.
            # We try to use a fallback that might be on the system or just use fpdf's core unicode handling
            
            try:
                # Try to load a common Windows Unicode font if on Windows
                # fpdf2 supports 'Arial' as a core font but it's not unicode.
                # However, fpdf2 can handle unicode if we don't specify a non-unicode font.
                pdf.set_font("Helvetica", size=12)
            except:
                pdf.set_font("Arial", size=12)

            # Metadata section removed as requested
            pdf.ln(5)
            
            # Content
            pdf.set_font(size=12)
            
            # Handle long content with better formatting
            # Split by paragraphs
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    # Handle bullet points specially
                    if para.strip().startswith('•') or para.strip().startswith('-'):
                        pdf.set_font(size=11)
                        lines = para.split('\n')
                        for line in lines:
                            if line.strip():
                                pdf.multi_cell(0, 6, line.strip())
                        pdf.ln(3)
                        pdf.set_font(size=12)
                    else:
                        pdf.multi_cell(0, 7, para.strip())
                        pdf.ln(5)
            
            # Footer
            if template in ["professional", "executive"]:
                pdf.set_y(-20)
                pdf.set_font(size=8)
                pdf.set_text_color(128, 128, 128)
                pdf.cell(0, 10, f"Page {pdf.page_no()}", align="C")
            
            # Save
            output_dir = "static/downloads"
            os.makedirs(output_dir, exist_ok=True)
            
            timestamp = int(datetime.now().timestamp())
            filename = f"document_{timestamp}.pdf"
            filepath = os.path.join(output_dir, filename)
            
            # fpdf2's output() supports unicode by default
            pdf.output(filepath)
            
            public_url = f"/static/downloads/{filename}"
            
            print(f"✅ PDF generated: {filename}")
            
            return {
                "download_url": public_url,
                "output": f"✅ PDF generated successfully!\n\nDownload: {public_url}",
                "filename": filename,
                "format": "pdf",
                "file_path": filepath
            }
            
        except Exception as e:
            print(f"❌ PDF Error: {str(e)}")
            return {
                "error": str(e),
                "output": f"❌ Failed to generate PDF: {str(e)}"
            }
    
    def _generate_docx(
        self,
        content: str,
        title: str,
        template: str,
        metadata: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Generate DOCX document."""
        if not Document:
            return {"error": "python-docx library not installed", "output": "❌ DOCX Error: Library missing"}
        
        try:
            doc = Document()
            
            # Title
            title_para = doc.add_heading(title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata section removed as requested
            doc.add_paragraph()  # Spacer
            
            # Content
            paragraphs = content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    # Check if bullet points
                    if para_text.strip().startswith('•') or para_text.strip().startswith('-'):
                        lines = para_text.split('\n')
                        for line in lines:
                            if line.strip():
                                # Remove existing bullet and add as list item
                                text = line.strip().lstrip('•-').strip()
                                doc.add_paragraph(text, style='List Bullet')
                    else:
                        doc.add_paragraph(para_text.strip())
            
            # Save
            output_dir = "static/downloads"
            os.makedirs(output_dir, exist_ok=True)
            
            timestamp = int(datetime.now().timestamp())
            filename = f"document_{timestamp}.docx"
            filepath = os.path.join(output_dir, filename)
            
            doc.save(filepath)
            
            public_url = f"/static/downloads/{filename}"
            
            print(f"✅ DOCX generated: {filename}")
            
            return {
                "download_url": public_url,
                "output": f"✅ DOCX generated successfully!\n\nDownload: {public_url}",
                "filename": filename,
                "format": "docx",
                "file_path": filepath
            }
            
        except Exception as e:
            print(f"❌ DOCX Error: {str(e)}")
            return {
                "error": str(e),
                "output": f"❌ Failed to generate DOCX: {str(e)}"
            }
    
    def _generate_txt(
        self,
        content: str,
        title: str,
        metadata: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Generate plain text document."""
        try:
            # Format text document
            output_text = f"{title}\n"
            output_text += "=" * len(title) + "\n\n"
            # Metadata section removed as requested
            output_text += "\n"
            output_text += content
            
            # Save
            output_dir = "static/downloads"
            os.makedirs(output_dir, exist_ok=True)
            
            timestamp = int(datetime.now().timestamp())
            filename = f"document_{timestamp}.txt"
            filepath = os.path.join(output_dir, filename)
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(output_text)
            
            public_url = f"/static/downloads/{filename}"
            
            print(f"✅ TXT generated: {filename}")
            
            return {
                "download_url": public_url,
                "output": f"✅ TXT generated successfully!\n\nDownload: {public_url}",
                "filename": filename,
                "format": "txt",
                "file_path": filepath
            }
            
        except Exception as e:
            print(f"❌ TXT Error: {str(e)}")
            return {
                "error": str(e),
                "output": f"❌ Failed to generate TXT: {str(e)}"
            }
    
    def _generate_markdown(
        self,
        content: str,
        title: str,
        metadata: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Generate Markdown document."""
        try:
            # Format markdown
            output_md = f"# {title}\n\n"
            # Metadata section removed as requested
            output_md += "---\n\n"
            output_md += content
            
            # Save
            output_dir = "static/downloads"
            os.makedirs(output_dir, exist_ok=True)
            
            timestamp = int(datetime.now().timestamp())
            filename = f"document_{timestamp}.md"
            filepath = os.path.join(output_dir, filename)
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(output_md)
            
            public_url = f"/static/downloads/{filename}"
            
            print(f"✅ Markdown generated: {filename}")
            
            return {
                "download_url": public_url,
                "output": f"✅ Markdown generated successfully!\n\nDownload: {public_url}",
                "filename": filename,
                "format": "md",
                "file_path": filepath
            }
            
        except Exception as e:
            print(f"❌ Markdown Error: {str(e)}")
            return {
                "error": str(e),
                "output": f"❌ Failed to generate Markdown: {str(e)}"
            }
